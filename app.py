from __future__ import annotations

import os
import re
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from flask import Flask, render_template, request, send_file
from flask_cors import CORS
from dotenv import load_dotenv
from werkzeug.utils import secure_filename

def _load_allowed_extensions() -> set[str]:
    raw = os.environ.get("ALLOWED_EXTENSIONS", ".docx")
    extensions: set[str] = set()
    for item in raw.split(","):
        ext = item.strip().lower()
        if not ext:
            continue
        if not ext.startswith("."):
            ext = f".{ext}"
        extensions.add(ext)
    return extensions


def _load_signature_lines() -> list[str]:
    raw = os.environ.get(
        "SIGNATURE_LINES", "D. Stephen Scherer|Bar No. 36003"
    )
    return [line.strip() for line in raw.split("|") if line.strip()]


def _load_frontend_origins() -> list[str] | str:
    raw = os.environ.get("FRONTEND_ORIGINS", "*").strip()
    if raw == "*":
        return "*"
    return [origin.strip() for origin in raw.split(",") if origin.strip()]


load_dotenv()

ALLOWED_EXTENSIONS = _load_allowed_extensions()
SIGNATURE_LINES = _load_signature_lines()
FONT_NAME = os.environ.get("FONT_NAME", "Script MT Bold")

app = Flask(__name__)
CORS(app, resources={r"/upload": {"origins": _load_frontend_origins()}})


def is_allowed(filename: str) -> bool:
    return Path(filename).suffix.lower() in ALLOWED_EXTENSIONS


def apply_signature(doc: Document) -> None:
    paragraph = doc.add_paragraph()

    # Ensure run properties + font elements exist before setting names.
    first_run = paragraph.add_run(SIGNATURE_LINES[0])
    first_run.font.name = FONT_NAME
    first_run.bold = True
    first_rpr = first_run._element.get_or_add_rPr()
    first_rfonts = first_rpr.get_or_add_rFonts()
    first_rfonts.set(qn("w:ascii"), FONT_NAME)
    first_rfonts.set(qn("w:hAnsi"), FONT_NAME)
    first_rfonts.set(qn("w:cs"), FONT_NAME)
    first_rfonts.set(qn("w:eastAsia"), FONT_NAME)

    first_run.add_break()

    # Apply the same font settings to the second line.
    second_run = paragraph.add_run(SIGNATURE_LINES[1])
    second_run.font.name = FONT_NAME
    second_run.bold = True
    second_rpr = second_run._element.get_or_add_rPr()
    second_rfonts = second_rpr.get_or_add_rFonts()
    second_rfonts.set(qn("w:ascii"), FONT_NAME)
    second_rfonts.set(qn("w:hAnsi"), FONT_NAME)
    second_rfonts.set(qn("w:cs"), FONT_NAME)
    second_rfonts.set(qn("w:eastAsia"), FONT_NAME)


@app.route("/", methods=["GET"])
def index() -> str:
    return render_template("index.html")


@app.route("/upload", methods=["POST"])
def upload():
    if "document" not in request.files:
        return "Missing file", 400

    file = request.files["document"]
    if file.filename == "":
        return "Missing file", 400

    original_name = secure_filename(file.filename)
    if not is_allowed(original_name):
        return "Only .docx files are supported", 400

    with tempfile.TemporaryDirectory() as tmp_dir:
        input_path = Path(tmp_dir) / original_name
        file.save(input_path)

        doc = Document(input_path)
        apply_signature(doc)

        reviewed_name = re.sub(
            r"unreviewed", "Reviewed", original_name, flags=re.IGNORECASE
        )
        output_path = Path(tmp_dir) / reviewed_name
        doc.save(output_path)

        return send_file(
            output_path,
            as_attachment=True,
            download_name=reviewed_name,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )


if __name__ == "__main__":
    port = int(os.environ.get("PORT", "5000"))
    app.run(host="0.0.0.0", port=port, debug=True)
